"""
Create a changelog fron conventional commit messages according to configuration

Format of the commit message:
<type>[!][(<optional scope>)]: <description>

[optional body]

[optional footer(s)]

With type one of:
•	feat: new feature
•	fix: bug fix on feature already in production
•	miss: something miss (or a bug) in a feature that was not already delivered in production (so will normally be hidden in change log)
•	perf: improves in performance
•	refactor: refactoring production code
•	docs: changes to documentation
•	style: formatting, spelling correction ..., etc; no code change
•	test: adding missing tests, refactoring tests; no production code change
•	chore: any other tasks like updating build,... etc; no production code change

And finally to revert a previous commit use revert as type and add the "Refs:" (sha) of the corresponding commits in the footer:
    revert: let us never again speak of the noodle incident

    Refs: 676104e, a215868

The  ! means a breaking change, we can also add in the footer
BREAKING CHANGE: <description>
footers other than BREAKING CHANGE: <description> may be provided and follow a convention similar to git trailer format.
"""
from git import Repo, cmd
from docx import Document
from gitdb.exc import BadName
import re
import sys
import argparse
import logging

types=["feat", "fix", "miss", "perf", "refactor", "docs", "style", "test", "chore"]
revert="revert"
name=dict(zip(types,map(str.capitalize,types)))  
name.update({"feat":"Feature", "fix": "Bug fix"})

parser = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
parser.add_argument("--version", help='version number (you can use VersionUpdate.exe to retrieve it)')
parser.add_argument("--types", help='types to show in changelog (default feat and fix)', choices=types, nargs="+", action='extend')
parser.add_argument("--start", help='overwrite the sha of the commit to start (else it is in the comment section of the docx)')
parser.add_argument("--end", help='overwrite the sha of the last commit (else HEAD)', default="")
parser.add_argument("docx", help='path to the docx')
parser.add_argument("--log", help='log level (default INFO)', default="INFO", choices=['DEBUG', 'INFO', 'WARN', 'ERROR'])
args = parser.parse_args()

numeric_level = getattr(logging, args.log.upper(), None)

if not isinstance(numeric_level, int):
    raise ValueError('Invalid log level: %s' % args.log)
logging.basicConfig(level=numeric_level,  format='%(levelname)s: %(message)s')

if not args.types:
    args.types = types[:2] # default is feat and fix
args.types.append(revert)
types.append(revert)

repo=Repo()
g=cmd.Git()

end_sha = repo.commit(args.end).hexsha
# if version is not passed, use the last tag name
version_name = args.version or g.describe("--abbrev=0","--tags", end_sha)

try:
    document = Document(args.docx)
except Exception as e:
    logging.error(e)
    sys.exit(1)

hexsha = args.start
if not hexsha:
    hexsha = document.core_properties.comments

# check existing commit in docx
if not hexsha:
    # else use the commit of the last tag
    tag=g.describe("--abbrev=0","--tags",repo.commit(args.end).parents[0].hexsha) 
    t=repo.tags[tag]
    hexsha=t.commit.parents[0].hexsha
    logging.warning(f"no version in docx use {tag}")

try:
    repo.commit(hexsha)
except BadName as e:
    logging.error(e)
    sys.exit(1)

if hexsha == end_sha:
    logging.error("No changes")
    sys.exit(2)

# check all changes
changes =[]
for c in repo.iter_commits(hexsha+".."+end_sha):
    msg = c.message.split("\n")
    # remove the bad comment from Azure
    msg[0] = re.sub(r"Merged PR \d+:\s*", "", msg[0])
    conv = re.search(r"(?P<type>"+"|".join(types)+r")(?P<break>!?)(\((?P<scope>\w+)\))?:\s*(?P<subject>.*)", msg[0])
    if conv:
        t = conv.group("type")
        if t in args.types:
            if t=="revert":
                logging.warning(f"{c.hexsha}: revert not done, should be done manually") 
            else:
                data = conv.groupdict()
                data["type"] = name[t]
                data["body"] = []
                data["footer"] = []
                # try to split body and footer
                i=[x[0] for x in enumerate(msg) if x[1]==""] 
                if len(i)>2:
                    data["body"] = [m for m in msg[i[0]+1:i[1]] if m != msg[0]]
                    data["footer"] = [m for m in msg[i[1]+1:i[-1]] if m]
                    if len(i)>3:
                        logging.warning(f"{c.hexsha}: badly formatted commit message")
                elif len(i)>1:
                    if msg[i[0]+1].startswith("Related work items:") or msg[i[0]+1].startswith("BREAKING CHANGE:"):
                        data["footer"] = msg[i[0]+1:i[1]]
                    else:
                        data["body"] = [m for m in msg[i[0]+1:i[1]] if m != msg[0]]

                for m in data["footer"]:
                    if not re.match(r"[\w ]+:", m):
                        logging.warning(f"{c.hexsha}: badly formatted footer")
                        break
                changes.append(data)
        else:
            logging.debug(f'ignore type {t} on {c.hexsha}')
    else:
        logging.warning(f'no type found on {c.hexsha}')

# add text a the begining of the document if there is content
if document.paragraphs: 
    addp = document.paragraphs[0].insert_paragraph_before
else:
    addp = document.add_paragraph

# file the docx with release information
p=addp(version_name, style="Mouvent Title")
for c in changes:
    if c["scope"]:
        c["type"] = "%(type)s(%(scope)s)" % c
    addp("%(type)s: %(subject)s" % c, style=(c["break"] and "Breaking" or "List Bullet"))
    if c["body"]:
        addp("\n".join(c["body"]), style="Body")
    for m in c["footer"]:
        if m.startswith("BREAKING CHANGE:"):
            addp(m, style="breaking change")

document.core_properties.comments = end_sha
document.save(args.docx)
